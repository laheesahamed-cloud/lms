from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "lms-beginner-owner-guide.docx"


COLORS = {
    "ink": "1F2937",
    "muted": "52616B",
    "blue": "2563A8",
    "dark_blue": "17324D",
    "light_blue": "EAF3FB",
    "gray_fill": "F4F6F8",
    "yellow_fill": "FFF7D6",
    "green_fill": "EAF7EF",
    "red_fill": "FDECEC",
    "border": "D4DCE3",
}


def run(cmd: list[str]) -> str:
    try:
        return subprocess.check_output(cmd, cwd=ROOT, text=True, stderr=subprocess.DEVNULL).strip()
    except Exception:
        return ""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["border"], size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in_inches):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def set_table_style(table, header_fill: str = COLORS["gray_fill"]) -> None:
    table.style = "Table Grid"
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_end)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 24, COLORS["dark_blue"], 0, 8),
        ("Subtitle", 12, COLORS["muted"], 0, 18),
        ("Heading 1", 16, COLORS["blue"], 16, 7),
        ("Heading 2", 13, COLORS["blue"], 12, 5),
        ("Heading 3", 11.5, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
        if name == "Title":
            style.font.bold = True
        if name.startswith("Heading"):
            style.font.bold = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.12

    if "CodeText" not in styles:
        code = styles.add_style("CodeText", 1)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(9)
        code.font.color.rgb = RGBColor.from_string("111827")
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.left_indent = Inches(0.16)
        code.paragraph_format.right_indent = Inches(0.16)

    if "SmallMuted" not in styles:
        muted = styles.add_style("SmallMuted", 1)
        muted.font.name = "Calibri"
        muted.font.size = Pt(9)
        muted.font.color.rgb = RGBColor.from_string(COLORS["muted"])
        muted.paragraph_format.space_after = Pt(3)


def p(doc: Document, text: str = "", style: str | None = None, bold: bool = False, italic: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
    return paragraph


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def bullet(doc: Document, text: str) -> None:
    p(doc, text, style="List Bullet")


def number(doc: Document, text: str) -> None:
    p(doc, text, style="List Number")


def code(doc: Document, text: str) -> None:
    paragraph = p(doc, text, style="CodeText")
    set_paragraph_shading(paragraph, COLORS["gray_fill"])


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def callout(doc: Document, title: str, body: str, fill: str = COLORS["light_blue"]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_border(cell, COLORS["border"])
    set_cell_margins(cell, 120, 160, 120, 160)
    para = cell.paragraphs[0]
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    para.add_run(" " + body)
    p(doc, "", style="SmallMuted")


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.9, 4.6), header: tuple[str, str] | None = None) -> None:
    count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=count, cols=2)
    table.autofit = False
    set_table_width(table, list(widths))
    start = 0
    if header:
        table.cell(0, 0).text = header[0]
        table.cell(0, 1).text = header[1]
        start = 1
    for idx, (left, right) in enumerate(rows, start=start):
        table.cell(idx, 0).text = left
        table.cell(idx, 1).text = right
    set_table_style(table)
    p(doc, "", style="SmallMuted")


def add_three_col_table(doc: Document, header: tuple[str, str, str], rows: list[tuple[str, str, str]], widths=(1.55, 2.15, 2.8)) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.autofit = False
    set_table_width(table, list(widths))
    for i, text in enumerate(header):
        table.cell(0, i).text = text
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    set_table_style(table)
    p(doc, "", style="SmallMuted")


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("xyndrome LMS Owner Guide | Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    add_field(paragraph, "PAGE")
    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_cover(doc: Document, branch: str, remote: str) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("xyndrome LMS Beginner Owner Guide")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("A plain-English manual for opening, understanding, changing, and launching your app")
    p(doc, "", style="SmallMuted")

    callout(
        doc,
        "Read this first.",
        "You do not need to become a software engineer to own this project. "
        "This guide explains the project like a map: where it lives, what each major folder means, "
        "how to open it locally, what Git main and branches mean, and what still matters before launch.",
        COLORS["green_fill"],
    )

    add_kv_table(
        doc,
        [
            ("Project folder", str(ROOT)),
            ("Local website", "http://localhost/lms/"),
            ("Login page", "http://localhost/lms/auth/login"),
            ("Backend API", "http://localhost:3000/api"),
            ("Current Git branch", branch or "unknown"),
            ("GitHub remote", remote or "not detected"),
            ("Generated on", date.today().isoformat()),
        ],
        header=("Item", "Value"),
    )

    h2(doc, "How to use this PDF")
    number(doc, "Start with Part 1 if you only want the simple picture.")
    number(doc, "Use Part 2 when you need to open the LMS locally.")
    number(doc, "Use Part 3 when you ask Codex or ChatGPT to change the app.")
    number(doc, "Use Part 4 before launch.")
    number(doc, "Use the appendices as dictionaries for files, commands, and Git words.")


def add_toc(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Contents")
    for item in [
        "Part 1 - The simple picture",
        "Part 2 - How to open the LMS locally",
        "Part 3 - Project folders and what they mean",
        "Part 4 - Git, main branch, and sub-branches",
        "Part 5 - How changes should be made safely",
        "Part 6 - Launch readiness",
        "Part 7 - Native app, API IPs, and local testing",
        "Part 8 - Troubleshooting",
        "Appendix A - File and folder dictionary",
        "Appendix B - Commands cheat sheet",
        "Appendix C - Environment variables in plain English",
        "Appendix D - What to ask Codex",
        "Appendix E - Learning path for the next 30 days",
    ]:
        bullet(doc, item)
    callout(
        doc,
        "Note.",
        "This table of contents is a reading map, not a clickable Word-generated TOC. "
        "The PDF itself is meant to be opened and skimmed visually.",
        COLORS["yellow_fill"],
    )


def part_simple_picture(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 1 - The Simple Picture")
    p(
        doc,
        "Your LMS is not one single magic file. It is a small system made of a few parts that talk to each other. "
        "When something breaks, the goal is to identify which part is failing instead of feeling like the whole app is broken.",
    )
    add_three_col_table(
        doc,
        ("Part", "Plain-English Meaning", "Real Project Location"),
        [
            ("Frontend", "The screens people see and click: website, login, student dashboard, admin pages, quiz pages.", "frontend/src/"),
            ("Backend/API", "The server brain. It receives requests, checks users, saves quiz results, talks to the database.", "backend/src/"),
            ("Database", "The stored data: users, courses, lessons, questions, quizzes, payments, attempts.", "MySQL through XAMPP or production DB"),
            ("XAMPP/Apache", "The local web server that lets the built website open at localhost/lms.", "/Applications/XAMPP/"),
            ("Native app", "Android/iOS wrappers around the frontend, using Capacitor.", "frontend/android/ and frontend/ios/"),
            ("Desktop app", "Electron wrapper for desktop builds.", "desktop/"),
            ("Git", "The project history and safety system.", ".git/ and GitHub"),
        ],
    )

    h2(doc, "The request flow")
    p(doc, "Most screens follow this flow:")
    number(doc, "User opens a page in the browser or native app.")
    number(doc, "The React frontend displays the screen.")
    number(doc, "If the page needs data, frontend calls the backend API.")
    number(doc, "Backend checks permissions and talks to MySQL.")
    number(doc, "Backend sends data back.")
    number(doc, "Frontend shows the result.")
    code(doc, "Browser or mobile app -> React frontend -> /api backend -> MySQL database -> response back to screen")

    h2(doc, "Why local development can still work after adding a real domain")
    p(
        doc,
        "Files such as sitemap.xml, robots.txt, canonical URLs, and social preview tags can point to your real public domain. "
        "That does not stop local development. Local development uses localhost and local API settings. "
        "SEO metadata is mostly for Google and social previews after deployment.",
    )
    add_kv_table(
        doc,
        [
            ("Local app URL", "http://localhost/lms/"),
            ("Local API URL", "http://localhost:3000/api"),
            ("Production public URL", "https://xyndrome.lk or your chosen final URL"),
            ("Production API URL", "https://xyndrome.lk/api"),
        ],
        header=("Type", "Example"),
    )

    h2(doc, "What vibe coding means for ownership")
    p(
        doc,
        "Vibe coding is useful because you can describe what you want and let an AI agent help implement it. "
        "The important upgrade is learning enough of the map to ask for safer changes and to know what must be tested. "
        "You do not need to memorize every file. You need to know which area the change belongs to and what result to verify.",
    )
    callout(
        doc,
        "Owner mindset.",
        "You are allowed to say: 'I do not know the file names. Please inspect the app, find the right files, make the change, run tests, and explain what changed.' That is a good prompt.",
        COLORS["green_fill"],
    )


def part_open_locally(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 2 - How To Open The LMS Locally")
    p(
        doc,
        "Local means the app runs on your own computer instead of the public internet. "
        "Your local app is stored inside the XAMPP htdocs folder.",
    )
    code(doc, str(ROOT))

    h2(doc, "Daily local startup")
    number(doc, "Open XAMPP.")
    number(doc, "Start Apache.")
    number(doc, "Start MySQL.")
    number(doc, "Open Terminal inside the project folder.")
    number(doc, "Start the backend API.")
    number(doc, "Open the login page in the browser.")
    code(doc, "cd /Applications/XAMPP/xamppfiles/htdocs/lms\nnpm run start:api:bg\nnpm run status:api\nopen http://localhost/lms/auth/login")

    h2(doc, "The most important local URLs")
    add_kv_table(
        doc,
        [
            ("Public website", "http://localhost/lms/"),
            ("Login", "http://localhost/lms/auth/login"),
            ("Student dashboard", "http://localhost/lms/app/dashboard"),
            ("Admin dashboard", "http://localhost/lms/admin/dashboard"),
            ("API health", "http://localhost:3000/api/health"),
            ("API ready check", "http://localhost:3000/api/health/ready"),
        ],
        header=("Area", "URL"),
    )

    h2(doc, "If the login page says Network Error")
    p(doc, "This usually means the frontend opened, but the backend API is not running or cannot reach the database.")
    number(doc, "Check whether the API is running.")
    code(doc, "npm run status:api")
    number(doc, "Check API health.")
    code(doc, "curl http://localhost:3000/api/health")
    number(doc, "If needed, restart the API.")
    code(doc, "npm run stop:api:bg\nnpm run start:api:bg")
    number(doc, "If health works but ready fails, check XAMPP MySQL and database settings.")

    h2(doc, "Why /api alone can show 404")
    p(
        doc,
        "The backend does not have a plain index page at /api. That is normal. "
        "Use a real endpoint like /api/health, /api/auth/login, or /api/results.",
    )

    h2(doc, "When frontend changes do not appear")
    p(
        doc,
        "The Apache version uses the built frontend bundle from frontend/dist. "
        "After changing frontend files, rebuild the frontend before testing through localhost/lms.",
    )
    code(doc, "npm run build:frontend")
    callout(
        doc,
        "Simple rule.",
        "If you changed what the user sees and localhost/lms still looks old, rebuild the frontend.",
        COLORS["yellow_fill"],
    )


def part_folders(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 3 - Project Folders And What They Mean")
    p(
        doc,
        "You do not need to open every folder. Think of the project like a building. "
        "Some rooms are where the app is built. Some are storage. Some are generated and should not be edited by hand.",
    )
    add_three_col_table(
        doc,
        ("Folder/File", "What It Means", "Should You Edit It?"),
        [
            ("frontend/src/surfaces/website/", "Landing page, login, register, legal pages, public AI generator.", "Only when changing website/auth/public pages."),
            ("frontend/src/surfaces/app/", "Student app: dashboard, courses, lessons, quizzes, results, notes, billing.", "Only when changing student experience."),
            ("frontend/src/surfaces/admin/", "Admin dashboard and management pages.", "Only when changing admin experience."),
            ("frontend/src/shared/", "Shared UI, API clients, auth store, layout, platform helpers, styles.", "Careful. Changes can affect many pages."),
            ("frontend/public/", "Public images, icons, robots.txt, sitemap.xml.", "Yes for assets and SEO files."),
            ("frontend/dist/", "Built web output generated by Vite.", "No. Rebuild instead."),
            ("frontend/dist-capacitor/", "Built native output generated for Capacitor.", "No. Rebuild/sync instead."),
            ("backend/src/modules/", "Backend features: auth, courses, questions, quizzes, payments, settings.", "Yes, but test carefully."),
            ("backend/.env", "Local backend secrets and database settings.", "Edit locally, never share publicly."),
            ("backend/.env.example", "Example backend environment template.", "Safe template only. No real secrets."),
            ("database/", "SQL dumps and migrations.", "Careful. Database changes need backup."),
            ("docs/", "Project documentation and guides.", "Yes."),
            ("scripts/", "Build, QA, migration, and helper scripts.", "Only when changing workflow."),
            ("desktop/", "Electron desktop wrapper.", "Only for desktop app changes."),
            ("frontend/android/ and frontend/ios/", "Native app projects generated/managed by Capacitor.", "Usually through sync, not manual edits."),
            ("node_modules/", "Installed packages.", "No. Do not edit."),
            ("backups/", "Saved project/database copies.", "Do not casually delete."),
        ],
    )

    h2(doc, "Where common changes usually live")
    add_three_col_table(
        doc,
        ("Change You Want", "Likely Area", "Example"),
        [
            ("Landing page text/design", "frontend/src/surfaces/website/pages/", "LandingPage.jsx"),
            ("Login/register UI", "frontend/src/surfaces/website/auth/", "LoginPage.jsx, RegisterPage.jsx"),
            ("Student dashboard", "frontend/src/surfaces/app/student/dashboard/", "StudentDashboardPage.jsx"),
            ("Quiz taking screen", "frontend/src/surfaces/app/student/quizzes/", "TakeQuizPage.jsx"),
            ("Quiz result/review", "frontend/src/surfaces/app/student/results/", "ResultPage.jsx, ReviewPage.jsx"),
            ("Admin questions", "frontend/src/surfaces/admin/pages/questions/", "QuestionsPage.jsx"),
            ("API request code", "frontend/src/shared/api/", "auth.api.js, aiNotes.api.js, etc."),
            ("Dark mode colors", "frontend/src/shared/styles/", "CSS files and variables"),
            ("Backend quiz logic", "backend/src/modules/quiz-attempts/", "controller/service files"),
            ("Payment settings", "backend/src/modules/subscriptions/ and settings", "PayHere flows"),
            ("Email/password reset", "backend/src/modules/auth/ and settings", "SMTP settings"),
            ("API IP/domain config", "frontend env files and shared platform config", ".env.capacitor, config files"),
        ],
    )

    h2(doc, "Files recently important in your project")
    p(doc, "These are areas that came up during recent work and launch cleanup.")
    bullet(doc, "frontend/src/surfaces/app/student/quizzes/TakeQuizPage.jsx - quiz-taking and finish animation behavior.")
    bullet(doc, "frontend/src/shared/seo/PageMeta.jsx - page title, description, canonical, social preview metadata.")
    bullet(doc, "frontend/public/robots.txt and frontend/public/sitemap.xml - Google crawl guidance.")
    bullet(doc, "frontend/src/surfaces/website/pages/RefundPolicyPage.jsx and CookiePolicyPage.jsx - legal pages.")
    bullet(doc, "frontend/src/shared/platform/config.js - platform/API fallback configuration.")
    bullet(doc, "scripts/build/sync-root-index.mjs - copies built frontend files to the root served by Apache.")


def part_git(doc: Document, branch: str, remote: str) -> None:
    page_break(doc)
    h1(doc, "Part 4 - Git, Main Branch, And Sub-Branches")
    p(
        doc,
        "Git is the save-history system for your app. GitHub is the online place where that history can be stored. "
        "A branch is like a separate working lane. You can try changes on a branch without immediately changing the stable main lane.",
    )
    callout(
        doc,
        "Important correction.",
        "A branch is not really a folder inside main. It is more like a named timeline. People may say 'sub-branch', but Git usually just says branch.",
        COLORS["light_blue"],
    )

    add_kv_table(
        doc,
        [
            ("Repository", "The full project with history."),
            ("main branch", "Usually the stable production branch. Treat it as the clean copy."),
            ("feature branch", "A separate branch for one fix or feature."),
            ("current branch", branch or "unknown"),
            ("commit", "A saved checkpoint with a message."),
            ("push", "Upload your commits to GitHub."),
            ("pull", "Download latest commits from GitHub."),
            ("merge", "Bring one branch's changes into another branch."),
            ("pull request", "A GitHub review page before merging."),
            ("remote", remote or "The GitHub project URL, if connected."),
        ],
        header=("Git Word", "Plain-English Meaning"),
    )

    h2(doc, "Simple branch example")
    p(doc, "Imagine main is your clean notebook. A branch is a photocopy where you try a change.")
    number(doc, "main has the current stable app.")
    number(doc, "Create a branch called fix-login.")
    number(doc, "Change login code on fix-login.")
    number(doc, "Run tests.")
    number(doc, "Commit the change.")
    number(doc, "Push branch to GitHub.")
    number(doc, "Merge fix-login into main only when you are happy.")

    h2(doc, "Commands you can safely ask Codex to run")
    add_kv_table(
        doc,
        [
            ("Check current branch", "git branch --show-current"),
            ("Check changed files", "git status --short"),
            ("See a summary of edits", "git diff --stat"),
            ("Create a new branch", "git switch -c codex/short-change-name"),
            ("Switch branch", "git switch branch-name"),
            ("Save changes", "git add selected-files then git commit -m \"message\""),
            ("Upload branch", "git push -u origin branch-name"),
        ],
        header=("Goal", "Command"),
    )

    h2(doc, "Git safety rules")
    bullet(doc, "Do not commit real secrets such as database passwords, API keys, PayHere secrets, SMTP passwords, or private keys.")
    bullet(doc, "Do not run destructive commands unless you fully understand them.")
    bullet(doc, "Avoid git reset --hard unless you intentionally want to throw away local work.")
    bullet(doc, "Before major changes, ask Codex to check git status and explain what files are already modified.")
    bullet(doc, "Use one branch per meaningful change when possible.")
    bullet(doc, "Commit only after build/tests pass, unless the commit is intentionally a work-in-progress.")

    callout(
        doc,
        "Best prompt.",
        "Before editing, say: 'Check git status first. Do not overwrite unrelated changes. Make a safe branch if needed. Then fix the issue and run tests.'",
        COLORS["green_fill"],
    )


def part_safe_changes(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 5 - How Changes Should Be Made Safely")
    p(
        doc,
        "Every app change should follow a simple loop: understand, edit, build, test, explain. "
        "This protects you from accidental breakage, especially when using AI coding tools.",
    )
    h2(doc, "The safe change loop")
    number(doc, "Describe the change in plain English.")
    number(doc, "Ask Codex to inspect the relevant files before editing.")
    number(doc, "Ask Codex to make the smallest reasonable change.")
    number(doc, "Run frontend build if the user interface changed.")
    number(doc, "Run backend build or tests if API/database/auth/payment changed.")
    number(doc, "Open the app and manually test the exact user flow.")
    number(doc, "Ask for a summary of changed files and remaining risks.")

    h2(doc, "Which checks to run")
    add_three_col_table(
        doc,
        ("Type Of Change", "Minimum Check", "Why"),
        [
            ("Text, color, layout", "npm run build:frontend", "Confirms React/CSS build still works."),
            ("Login/register/auth", "npm test and manual login/register", "Auth bugs can lock users out."),
            ("Quiz submit/results", "npm test plus manual student quiz flow", "Scoring/results must be trusted."),
            ("Admin content", "npm run build:backend and manual admin CRUD", "Admin data changes affect courses/questions."),
            ("Payment", "Backend tests plus sandbox/live payment test", "Money flow must be tested outside code only."),
            ("Email reset", "SMTP test with real email", "Password reset depends on external email provider."),
            ("Native app", "npm run mobile:cap:sync then device test", "Native bundle can lag behind web code."),
            ("Production deploy", "npm test, npm run build, health checks, manual role checks", "Final safety gate."),
        ],
    )

    h2(doc, "What changed recently in plain English")
    bullet(doc, "Quiz finish animation was made more expressive and centered, with review transition behavior.")
    bullet(doc, "Native app support was synced after frontend changes.")
    bullet(doc, "Local native API IP was adjusted for your network, including 172.20.10.2 during testing.")
    bullet(doc, "Dark mode cards were made darker while keeping page backgrounds mostly unchanged.")
    bullet(doc, "SEO metadata, robots.txt, sitemap.xml, refund policy, and cookie policy were added.")
    bullet(doc, "A fake Google sign-in button was removed so users do not click an unfinished login option.")

    h2(doc, "Manual test checklist after a UI change")
    bullet(doc, "Open public website.")
    bullet(doc, "Open login page.")
    bullet(doc, "Login as student.")
    bullet(doc, "Open dashboard.")
    bullet(doc, "Start a quiz.")
    bullet(doc, "Finish quiz and check finish animation.")
    bullet(doc, "Confirm automatic review/result navigation works.")
    bullet(doc, "Switch light/dark mode if available.")
    bullet(doc, "Check on narrow/mobile width.")
    bullet(doc, "Logout and login as admin.")
    bullet(doc, "Check admin dashboard and one content page.")


def part_launch(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 6 - Launch Readiness")
    p(
        doc,
        "Launch is not only code. A real launch needs domain, HTTPS, production environment variables, database backups, payment setup, email setup, legal pages, monitoring, and manual testing.",
    )

    h2(doc, "Critical blockers before launch")
    add_three_col_table(
        doc,
        ("Blocker", "What Must Be True", "Who/Where"),
        [
            ("Domain", "A real domain is connected and production URLs point to it.", "Hosting/DNS and project env files."),
            ("HTTPS", "Website and API load over https, not plain http.", "Hosting/reverse proxy."),
            ("Production API", "Frontend calls the public production API, not localhost or LAN IP.", "Frontend env/build."),
            ("Database", "Production DB exists with correct user, password, schema, and backups.", "Hosting/MySQL."),
            ("Secrets", "JWT/session/settings keys are long and private.", "Backend environment."),
            ("Email", "Password reset sends real email successfully.", "SMTP settings."),
            ("Payment", "PayHere live payment and notify webhook are tested end to end.", "Admin settings/PayHere."),
            ("Legal", "Privacy, Terms, Refund, Cookie pages contain real business information.", "Website pages and legal review."),
            ("Monitoring", "Uptime/error alerts exist for API, database, payment, email, and disk space.", "Hosting/monitoring tool."),
            ("Manual QA", "Student, admin, logged-out, mobile, and payment flows are tested.", "You/Codex/testing person."),
        ],
    )

    h2(doc, "Files with domain placeholders")
    p(doc, "These files may contain production domain values. Replacing them does not stop local development.")
    bullet(doc, "frontend/index.html")
    bullet(doc, "frontend/public/robots.txt")
    bullet(doc, "frontend/public/sitemap.xml")
    bullet(doc, "backend/.env.example and production backend env")
    bullet(doc, "frontend/.env.example and production frontend env")
    bullet(doc, "frontend/.env.capacitor.example and native production env")

    h2(doc, "Production env values in plain English")
    add_kv_table(
        doc,
        [
            ("FRONTEND_URL", "Your public website URL."),
            ("FRONTEND_URLS", "Allowed browser origins that may call the API."),
            ("APP_PUBLIC_URL", "Public LMS app URL used in links."),
            ("API_PUBLIC_URL", "Public backend API URL."),
            ("ALLOW_LAN_ORIGINS", "Should be false in production."),
            ("SETTINGS_ENCRYPTION_KEY", "Long secret used to protect saved settings."),
            ("HEALTH_METRICS_TOKEN", "Secret token for protected metrics."),
            ("DB_HOST/DB_USER/DB_PASSWORD/DB_NAME", "Production database connection details."),
            ("VITE_API_BASE_URL", "Frontend's API target."),
            ("VITE_PUBLIC_WEBSITE_URL", "Public website URL for frontend metadata."),
        ],
        header=("Variable", "Meaning"),
    )

    h2(doc, "What not to launch with")
    bullet(doc, "Do not launch with localhost or LAN IP in production frontend API settings.")
    bullet(doc, "Do not launch with ALLOW_LAN_ORIGINS=true.")
    bullet(doc, "Do not launch with DB_USER=root or empty database password.")
    bullet(doc, "Do not launch if password reset email is untested.")
    bullet(doc, "Do not launch if PayHere live payment and notify webhook are untested.")
    bullet(doc, "Do not launch without a recent database backup.")
    bullet(doc, "Do not launch with example legal policy text only.")

    callout(
        doc,
        "Launch truth.",
        "A build passing means the code can compile. It does not prove payments, email, DNS, SSL, backups, or legal details are ready.",
        COLORS["yellow_fill"],
    )


def part_native(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 7 - Native App, API IPs, And Local Testing")
    p(
        doc,
        "The native app is built with Capacitor. Capacitor packages the React frontend into Android/iOS projects. "
        "The native app still needs to call the backend API, so API URLs matter a lot.",
    )

    h2(doc, "Why localhost can fail on a phone")
    p(
        doc,
        "On your Mac, localhost means your Mac. On a phone, localhost means the phone itself. "
        "So a phone usually cannot call http://localhost:3000/api unless the backend is running on the phone, which it is not.",
    )
    add_kv_table(
        doc,
        [
            ("Browser on Mac", "http://localhost:3000/api can work."),
            ("Phone on same Wi-Fi", "Use your Mac's LAN IP, such as http://172.20.10.2:3000/api when that is the current IP."),
            ("Production native app", "Use https://xyndrome.lk/api."),
        ],
        header=("Situation", "API URL"),
    )

    h2(doc, "Native sync command")
    p(doc, "After frontend/native changes, rebuild and sync Capacitor.")
    code(doc, "npm run mobile:cap:sync")
    p(doc, "For Android only:")
    code(doc, "npm run mobile:cap:sync:android")
    p(doc, "For iOS only:")
    code(doc, "npm run mobile:cap:sync:ios")

    h2(doc, "If the native app shows old UI")
    number(doc, "Rebuild/sync Capacitor.")
    number(doc, "Open Android Studio or Xcode again.")
    number(doc, "Clean/reinstall app if needed.")
    number(doc, "Confirm the device can reach the backend API URL.")

    h2(doc, "Production native rule")
    callout(
        doc,
        "Use HTTPS in production.",
        "A shipped native app should call the public HTTPS API. A local IP is only for testing on your own network.",
        COLORS["red_fill"],
    )


def part_troubleshooting(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Part 8 - Troubleshooting")
    p(doc, "Use this when something breaks. Start with the simplest checks before assuming the app is ruined.")
    add_three_col_table(
        doc,
        ("Symptom", "Most Likely Cause", "First Fix"),
        [
            ("Login says Network Error", "Backend API is stopped or unreachable.", "npm run start:api:bg then curl /api/health."),
            ("API health works but ready fails", "Database/MySQL problem.", "Start XAMPP MySQL and check backend .env."),
            ("/api shows 404", "No plain API index route.", "Use /api/health instead."),
            ("Frontend change not visible", "Built bundle is old.", "npm run build:frontend."),
            ("Native app old UI", "Capacitor not synced or app not reinstalled.", "npm run mobile:cap:sync."),
            ("Phone cannot reach API", "Using localhost instead of computer IP.", "Use current LAN IP or production HTTPS API."),
            ("Admin gets 403", "Role/permission problem.", "Check user role and backend permission mapping."),
            ("Results not showing", "Quiz attempts/student answers issue.", "Check quiz-attempts and results API."),
            ("Payment not activating", "PayHere notify/webhook not reaching backend.", "Check public notify URL and logs."),
            ("Password reset email not arriving", "SMTP setting or provider issue.", "Test SMTP settings and logs."),
        ],
    )

    h2(doc, "Log locations")
    bullet(doc, ".runtime/api.log - background API log.")
    bullet(doc, "Browser DevTools Network tab - failing frontend requests.")
    bullet(doc, "Backend terminal output - API errors during development.")
    bullet(doc, "Hosting provider logs - production backend errors.")

    h2(doc, "First three commands when confused")
    code(doc, "npm run status:api\ncurl http://localhost:3000/api/health\nnpm run build:frontend")
    p(doc, "If those are clean, then the problem is more specific and should be traced by page/API route.")


def appendix_files(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Appendix A - File And Folder Dictionary")
    p(doc, "This section is intentionally simple. Use it when someone mentions a file name and you want to know what it is.")
    add_three_col_table(
        doc,
        ("Name", "Meaning", "Beginner Warning"),
        [
            ("package.json", "A command/menu file for Node projects. It lists scripts like build and test.", "Edit only when changing commands or dependencies."),
            ("node_modules", "Installed packages downloaded by npm.", "Never edit manually."),
            ("src", "Source code written by developers/AI.", "This is where real code lives."),
            ("dist", "Built output generated from source.", "Do not edit; rebuild."),
            (".env", "Private local settings and secrets.", "Do not upload/share real secrets."),
            (".env.example", "A template showing what env variables exist.", "Safe only if it has no real secrets."),
            ("README.md", "Human-readable project notes.", "Good first file to read."),
            ("router.jsx", "Frontend route map: which URL opens which page.", "Wrong routes can break navigation."),
            ("AppFrame/AppShell", "Main app layout wrapper.", "Changes affect many pages."),
            ("BootLoader", "Loading screen when app starts.", "Affects perceived startup polish."),
            ("controller.ts", "Backend file that receives API requests.", "Usually paired with service.ts."),
            ("service.ts", "Backend file where business logic often lives.", "Changes can affect data behavior."),
            ("schema-sync.service.ts", "Database schema bootstrap/backfill logic.", "Be careful; database structure risk."),
            ("capacitor.config.ts", "Native app config.", "Important for mobile builds."),
            ("robots.txt", "Tells search engines what they may crawl.", "Use real domain at launch."),
            ("sitemap.xml", "List of important public URLs for search engines.", "Use real domain at launch."),
        ],
    )


def appendix_commands(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Appendix B - Commands Cheat Sheet")
    add_three_col_table(
        doc,
        ("Command", "Use It For", "When"),
        [
            ("npm run install:all", "Install frontend/backend dependencies.", "First setup or after package changes."),
            ("npm run start:api", "Start API in current terminal.", "Development when terminal stays open."),
            ("npm run start:api:bg", "Start API in background.", "Daily local use."),
            ("npm run status:api", "Check background API status.", "When login/API fails."),
            ("npm run stop:api:bg", "Stop background API.", "Restarting API."),
            ("npm run dev:frontend", "Start Vite dev server.", "Optional frontend development."),
            ("npm run build:frontend", "Build frontend.", "After UI changes."),
            ("npm run build:backend", "Build backend.", "After API/backend changes."),
            ("npm run build", "Build frontend and backend.", "Before committing/deploying."),
            ("npm test", "Run project QA/security checks.", "Before commit/launch."),
            ("npm run mobile:cap:sync", "Build and sync native bundle.", "After native/frontend changes."),
            ("npm run check:health", "Check local API health.", "When diagnosing API."),
            ("git status --short", "See changed files.", "Before and after edits."),
            ("git branch --show-current", "See branch name.", "Before committing."),
        ],
    )


def appendix_env(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Appendix C - Environment Variables In Plain English")
    p(
        doc,
        "Environment variables are settings outside the code. They tell the app where the database is, where the API is, and which secret keys to use. "
        "Production environment variables are usually set in the server or hosting dashboard, not typed into public GitHub files.",
    )
    h2(doc, "Backend variables")
    add_kv_table(
        doc,
        [
            ("NODE_ENV", "Use production on the live server."),
            ("PORT", "Backend server port, usually 3000 locally."),
            ("FRONTEND_URL/FRONTEND_URLS", "Allowed frontend domains that can call the API."),
            ("APP_PUBLIC_URL", "Public URL used in generated links."),
            ("API_PUBLIC_URL", "Public API URL."),
            ("BODY_LIMIT", "Max request body size."),
            ("ALLOW_LAN_ORIGINS", "Allows local network origins. False in production."),
            ("SETTINGS_ENCRYPTION_KEY", "Protects saved provider/payment/email settings."),
            ("HEALTH_METRICS_TOKEN", "Protects metrics endpoint."),
            ("OPENROUTER/GEMINI keys", "AI provider credentials."),
            ("VAPID/APNS/FCM", "Push notification credentials."),
            ("DB_*", "MySQL database connection."),
        ],
        header=("Backend Variable Group", "Meaning"),
    )
    h2(doc, "Frontend variables")
    add_kv_table(
        doc,
        [
            ("VITE_API_BASE_URL", "Main API URL the frontend calls."),
            ("VITE_API_BASE_URLS", "Fallback API URLs."),
            ("VITE_API_REQUEST_TIMEOUT_MS", "How long frontend waits for API response."),
            ("VITE_PWA_SW_URL/SCOPE", "Progressive Web App service worker paths."),
            ("VITE_LMS_BUILD_TARGET", "web, pwa, native, or desktop build target."),
            ("VITE_ENABLE_PWA", "Whether PWA features are enabled."),
            ("VITE_PUBLIC_WEBSITE_URL", "Public website URL used by metadata."),
            ("VITE_APP_ONLY_HOSTS", "Domains that should open app-only surface."),
        ],
        header=("Frontend Variable", "Meaning"),
    )
    callout(
        doc,
        "Secret rule.",
        "If a value can give access to money, user data, database, email, AI billing, or private admin functions, treat it as a secret.",
        COLORS["red_fill"],
    )


def appendix_codex(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Appendix D - What To Ask Codex")
    p(doc, "Good prompts tell Codex the result you want, the risk level, and what checks to run.")
    h2(doc, "General safe fix prompt")
    code(
        doc,
        "I do not know the file names. Please inspect the project, find the right files, fix this issue, do not overwrite unrelated changes, then run the relevant build/tests and explain what changed.",
    )
    h2(doc, "UI change prompt")
    code(
        doc,
        "Make this UI change: [describe]. Keep light/dark mode working. Check desktop and mobile layout. Run npm run build:frontend. Tell me which files changed.",
    )
    h2(doc, "Backend/API prompt")
    code(
        doc,
        "Fix this API/backend issue: [describe]. Check the route/controller/service. Run backend build and relevant tests. Explain any database risk.",
    )
    h2(doc, "Launch readiness prompt")
    code(
        doc,
        "Review launch readiness again. Check domain placeholders, env files, SEO files, legal pages, payment/email readiness, security, tests, and tell me what is still blocked.",
    )
    h2(doc, "Git prompt")
    code(
        doc,
        "Before editing, check git status and current branch. If needed, create a branch with codex/ prefix. After the fix, show changed files and ask before committing.",
    )


def appendix_learning(doc: Document) -> None:
    page_break(doc)
    h1(doc, "Appendix E - 30 Day Learning Path For Owning This App")
    p(doc, "This is not a computer science course. It is a practical path so you can talk to developers and AI tools with confidence.")
    add_three_col_table(
        doc,
        ("Days", "Focus", "Goal"),
        [
            ("1-3", "Open app locally", "Start XAMPP, API, login page, health check without panic."),
            ("4-6", "Folder map", "Know frontend, backend, database, docs, scripts, native folders."),
            ("7-9", "Git basics", "Understand main, branch, commit, push, pull, merge."),
            ("10-12", "Frontend basics", "Know where pages, shared components, styles, and API clients live."),
            ("13-15", "Backend basics", "Know controller, service, module, env, database connection."),
            ("16-18", "Debug basics", "Use API health, build output, logs, browser Network tab."),
            ("19-21", "Launch basics", "Understand domain, HTTPS, env vars, backups, monitoring."),
            ("22-24", "Payments/email", "Understand why PayHere and SMTP require real external testing."),
            ("25-27", "Native basics", "Understand localhost vs LAN IP vs production HTTPS API."),
            ("28-30", "Safe AI workflow", "Ask better prompts, demand tests, read summaries, avoid secrets."),
        ],
    )
    h2(doc, "The only things to memorize first")
    bullet(doc, "Project folder: /Applications/XAMPP/xamppfiles/htdocs/lms")
    bullet(doc, "Local login: http://localhost/lms/auth/login")
    bullet(doc, "API health: http://localhost:3000/api/health")
    bullet(doc, "Build frontend: npm run build:frontend")
    bullet(doc, "Run QA: npm test")
    bullet(doc, "Check branch: git branch --show-current")
    bullet(doc, "Check changed files: git status --short")


def build() -> None:
    branch = run(["git", "branch", "--show-current"])
    remote = run(["git", "remote", "get-url", "origin"])

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    add_footer(section)

    add_cover(doc, branch, remote)
    add_toc(doc)
    part_simple_picture(doc)
    part_open_locally(doc)
    part_folders(doc)
    part_git(doc, branch, remote)
    part_safe_changes(doc)
    part_launch(doc)
    part_native(doc)
    part_troubleshooting(doc)
    appendix_files(doc)
    appendix_commands(doc)
    appendix_env(doc)
    appendix_codex(doc)
    appendix_learning(doc)

    callout(
        doc,
        "Final reminder.",
        "You can own this app by using maps, checklists, backups, and safe Git branches. You do not need to memorize every file name. You need to ask for careful changes and verify the important flows.",
        COLORS["green_fill"],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
